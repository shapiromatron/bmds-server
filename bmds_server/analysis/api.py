from io import BytesIO, StringIO

import matplotlib.ticker as mtick
import pandas as pd
from bmds import plotting
from bmds.datasets.transforms import polyk
from django.core.exceptions import ValidationError
from docx import Document
from rest_framework import exceptions, mixins, viewsets
from rest_framework.decorators import action
from rest_framework.response import Response

from ..common import renderers
from ..common.renderers import BinaryFile
from ..common.serializers import UnusedSerializer
from ..common.task_cache import ReportStatus
from ..common.utils import get_bool
from ..common.validation import pydantic_validate
from . import models, schema, serializers, validators
from .reporting.cache import DocxReportCache, ExcelReportCache
from .reporting.docx import add_update_url


class AnalysisViewset(mixins.RetrieveModelMixin, viewsets.GenericViewSet):
    serializer_class = serializers.AnalysisSerializer
    queryset = models.Analysis.objects.all()

    @action(detail=False, url_path="default")
    def default(self, request, *args, **kwargs):
        data = models.Analysis().default_input()
        return Response(data)

    @action(
        detail=True,
        methods=("patch",),
        url_path="patch-inputs",
    )
    def patch_inputs(self, request, *args, **kwargs):
        """
        Validate input and if successful, patch inputs on server side.
        """
        instance = self.get_object()
        data = request.data.get("data")
        edit_key = request.data.get("editKey", "")
        partial = bool(request.data.get("partial", False))

        # permission check
        if edit_key != instance.password:
            raise exceptions.PermissionDenied()

        if not isinstance(data, dict):
            raise exceptions.ValidationError("A `data` object is required")

        try:
            validators.validate_input(data, partial=partial)
        except ValidationError as err:
            raise exceptions.ValidationError(err.message)

        instance.reset_execution()
        instance.inputs = data
        instance.save()

        serializer = self.get_serializer(instance)
        return Response(serializer.data)

    @action(detail=True, methods=("post",))
    def execute(self, request, *args, **kwargs):
        """
        Attempt to execute the model.
        """
        instance = self.get_object()

        # permissions check
        if instance.password != request.data.get("editKey", ""):
            raise exceptions.PermissionDenied()

        # preflight execution check
        if not instance.inputs_valid():
            return Response("Invalid inputs", status=400)
        elif instance.is_executing:
            return Response("Execution already started", status=400)

        # start analysis execution
        instance.reset_execution()
        instance.start_execute()

        instance.refresh_from_db()
        serializer = self.get_serializer(instance)
        return Response(serializer.data)

    @action(detail=True, methods=("post",), url_path="select-model")
    def select_model(self, request, *args, **kwargs):
        instance = self.get_object()

        # permissions check
        if instance.password != request.data.get("editKey", ""):
            raise exceptions.PermissionDenied()

        # validate data
        data = request.data.get("data")
        if not isinstance(data, dict):
            raise exceptions.ValidationError("A `data` object is required")

        selection = pydantic_validate(data, validators.AnalysisSelectedSchema)
        instance.update_selection(selection)

        # fetch from db and get the latest
        instance.refresh_from_db()
        serializer = self.get_serializer(instance)
        return Response(serializer.data)

    @action(detail=True, methods=("post",), url_path="execute-reset")
    def execute_reset(self, request, *args, **kwargs):
        """
        Attempt to execute the model.
        """
        instance = self.get_object()

        # permissions check
        if instance.password != request.data.get("editKey", ""):
            raise exceptions.PermissionDenied()

        # reset instance
        instance.reset_execution()
        instance.save()

        # fetch from db and get the latest
        instance.refresh_from_db()
        serializer = self.get_serializer(instance)
        return Response(serializer.data)

    @action(detail=True, renderer_classes=(renderers.XlsxRenderer,))
    def excel(self, request, *args, **kwargs):
        """
        Return Excel export of outputs for selected analysis
        """
        instance = self.get_object()
        cache = ExcelReportCache(analysis=instance)
        response = cache.request_content()
        if response.status is ReportStatus.COMPLETE:
            cache.delete()  # destroy from cache; request is now complete
            data = renderers.BinaryFile(data=response.content, filename=instance.slug)
            return Response(data)

        return Response(response.model_dump(), content_type="application/json")

    @action(detail=True, renderer_classes=(renderers.DocxRenderer,))
    def word(self, request, *args, **kwargs):
        """
        Return Word report for the selected analysis
        """
        instance: models.Analysis = self.get_object()
        uri = request.build_absolute_uri("/")[:-1]
        kwargs = {
            "dataset_format_long": get_bool(
                request.query_params.get("datasetFormatLong")
            ),
            "all_models": get_bool(request.query_params.get("allModels")),
            "bmd_cdf_table": get_bool(request.query_params.get("bmdCdfTable")),
        }
        cache = DocxReportCache(analysis=instance, uri=uri, **kwargs)
        response = cache.request_content()
        if response.status is ReportStatus.COMPLETE:
            cache.delete()  # destroy from cache; request is now complete
            edit = instance.password == request.query_params.get("editKey", "")
            data = (
                add_update_url(instance, response.content, uri)
                if edit
                else response.content
            )
            return Response(renderers.BinaryFile(data=data, filename=instance.slug))

        return Response(response.model_dump(), content_type="application/json")


class PolyKViewset(viewsets.GenericViewSet):
    queryset = models.Analysis.objects.none()
    serializer_class = UnusedSerializer

    def _run_analysis(self, request) -> tuple[pd.DataFrame, pd.DataFrame]:
        try:
            settings = pydantic_validate(request.data, schema.PolyKInput)
        except ValidationError as err:
            raise exceptions.ValidationError(err.message)
        return settings.calculate()

    def create(self, request, *args, **kwargs):
        df, df2 = self._run_analysis(request)
        return Response(
            {"df": df.to_dict(orient="list"), "df2": df2.to_dict(orient="list")}
        )

    def write_cell(self, cell, value):
        cell.text = str(value)

    def summary_figure(self):
        dfz = pd.read_csv(
            StringIO(
                "dose,day,has_tumor\n0,452,0\n0,535,0\n0,553,0\n0,570,0\n0,596,0\n0,603,0\n0,606,0\n0,607,0\n0,627,0\n0,628,0\n0,635,0\n0,635,0\n0,638,0\n0,664,0\n0,666,0\n0,674,0\n0,680,0\n0,684,0\n0,684,0\n0,719,0\n0,722,0\n0,729,0\n0,729,1\n0,733,1\n0,733,0\n0,733,0\n0,733,0\n0,733,0\n0,733,0\n0,733,0\n0,733,0\n0,733,0\n0,733,1\n0,733,0\n0,733,0\n0,733,0\n0,733,0\n0,733,0\n0,733,0\n0,733,0\n0,734,0\n0,734,0\n0,734,0\n0,734,0\n0,734,1\n0,734,0\n0,734,0\n0,734,0\n0,734,1\n0,734,1\n12.8,439,0\n12.8,523,0\n12.8,530,0\n12.8,546,0\n12.8,579,0\n12.8,582,0\n12.8,583,0\n12.8,596,0\n12.8,607,0\n12.8,614,0\n12.8,626,0\n12.8,635,0\n12.8,635,0\n12.8,638,1\n12.8,659,0\n12.8,667,0\n12.8,680,1\n12.8,688,0\n12.8,691,1\n12.8,691,0\n12.8,701,0\n12.8,708,0\n12.8,722,1\n12.8,733,0\n12.8,733,0\n12.8,733,1\n12.8,733,1\n12.8,733,0\n12.8,733,0\n12.8,733,0\n12.8,733,0\n12.8,733,0\n12.8,733,0\n12.8,733,1\n12.8,733,0\n12.8,733,0\n12.8,733,1\n12.8,733,1\n12.8,733,1\n12.8,733,0\n12.8,733,0\n12.8,733,0\n12.8,734,0\n12.8,734,1\n12.8,734,0\n12.8,734,0\n12.8,734,1\n12.8,734,0\n12.8,734,0\n12.8,734,0\n32,382,0\n32,439,0\n32,470,0\n32,491,0\n32,495,0\n32,502,0\n32,545,0\n32,551,0\n32,565,0\n32,567,1\n32,579,1\n32,596,1\n32,602,0\n32,606,1\n32,607,1\n32,607,0\n32,607,1\n32,620,1\n32,623,0\n32,629,0\n32,635,0\n32,649,0\n32,663,0\n32,666,0\n32,666,1\n32,679,1\n32,679,0\n32,680,1\n32,688,1\n32,688,0\n32,693,0\n32,701,1\n32,705,0\n32,713,1\n32,715,1\n32,717,1\n32,733,0\n32,733,1\n32,733,0\n32,733,1\n32,734,0\n32,734,1\n32,734,0\n32,734,0\n32,734,1\n32,734,1\n32,734,0\n32,734,1\n32,734,1\n32,734,1\n80,392,0\n80,422,0\n80,454,0\n80,523,0\n80,524,1\n80,567,1\n80,579,0\n80,587,1\n80,589,1\n80,595,0\n80,600,0\n80,603,0\n80,608,1\n80,620,1\n80,624,0\n80,624,1\n80,628,0\n80,635,0\n80,635,0\n80,637,0\n80,644,1\n80,644,0\n80,649,1\n80,649,1\n80,649,1\n80,656,1\n80,657,1\n80,659,1\n80,659,1\n80,663,0\n80,672,1\n80,677,1\n80,677,1\n80,680,0\n80,687,0\n80,701,0\n80,701,1\n80,733,1\n80,733,1\n80,733,1\n80,733,0\n80,733,1\n80,733,0\n80,733,1\n80,733,1\n80,733,1\n80,733,1\n80,733,1\n80,734,0\n80,734,0\n"
            )
        )
        df2z = polyk.adjust_n(dfz)
        df3 = polyk.summary_stats(df2z)
        units = "mg/kg/d"

        fig = plotting.create_empty_figure()
        ax = fig.gca()

        ax.set_xlabel(f"Dose ({units})" if units else "Dose")
        ax.set_ylabel("Proportion (%)")
        ax.margins(plotting.PLOT_MARGINS)
        ax.set_title("Adjusted Proportion vs Original Proportion")
        ax.plot(
            df3.dose,
            df3.proportion,
            "o-",
            color="blue",
            label="Original Proportion",
            markersize=8,
            markeredgewidth=1,
            markeredgecolor="white",
        )
        ax.plot(
            df3.dose,
            df3.adj_proportion,
            "^-",
            color="red",
            label="Adjusted Proportion",
            markersize=8,
            markeredgewidth=1,
            markeredgecolor="white",
        )
        ax.legend(**plotting.LEGEND_OPTS)
        ax.yaxis.set_major_formatter(mtick.PercentFormatter(1))
        plotting.close_figure(fig)
        return fig

    @action(detail=False, methods=["POST"], renderer_classes=(renderers.XlsxRenderer,))
    def excel(self, request, *args, **kwargs):
        df, df2 = self._run_analysis(request)
        f = BytesIO()
        # TODO - write real method here
        with pd.ExcelWriter(f) as writer:
            df.to_excel(writer, sheet_name="a", index=False)
            df2.to_excel(writer, sheet_name="b", index=False)
        data = BinaryFile(f, "demo")
        return Response(data)

    @action(detail=False, methods=["POST"], renderer_classes=(renderers.DocxRenderer,))
    def word(self, request, *args, **kwargs):
        df, df2 = self._run_analysis(request)
        f = BytesIO()
        doc = Document()
        doc.add_heading("Poly K Adjustment", 0)
        # Summary
        # Headers
        tbl = doc.add_table(df2.shape[0] + 1, 6)
        self.write_cell(tbl.cell(0, 0), "dose")
        self.write_cell(tbl.cell(0, 1), "n")
        self.write_cell(tbl.cell(0, 2), "adj_n")
        self.write_cell(tbl.cell(0, 3), "incidence")
        self.write_cell(tbl.cell(0, 4), "proportion")
        self.write_cell(tbl.cell(0, 5), "adj_proportion")

        for i, v in enumerate(
            zip(
                df2.dose,
                df2.n,
                df2.adj_n,
                df2.incidence,
                df2.proportion,
                df2.adj_proportion,
                strict=True,
            )
        ):
            # ew
            self.write_cell(tbl.cell(i + 1, 0), v[0])
            self.write_cell(tbl.cell(i + 1, 1), v[1])
            self.write_cell(tbl.cell(i + 1, 2), v[2])
            self.write_cell(tbl.cell(i + 1, 3), v[3])
            self.write_cell(tbl.cell(i + 1, 4), v[4])
            self.write_cell(tbl.cell(i + 1, 5), v[5])

        doc.add_heading("Plots", 1)
        # :(
        doc.add_paragraph(self.summary_figure())

        doc.add_heading("Table", 1)
        doc.add_heading("Data", 1)

        doc.save(f)
        data = BinaryFile(f, "demo")
        return Response(data)
