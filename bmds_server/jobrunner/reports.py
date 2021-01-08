from io import BytesIO
from typing import Dict, List

import bmds
import pandas as pd
from docx import Document

from .models import Job

# TODO - use contants in bmds library
# TODO - move this model to package from bmds-server to bmds?
bmrType = {
    1: "Abs. Dev",
    2: "Std. Dev",
    3: "Rel. Dev",
    4: "Point Estimate",
    5: "Hybrid-Extra Risk",
}
riskType = {1: "Extra Risk", 2: "Added Risk"}


class ReportEngine:
    def __init__(self, job: Job):
        self.job = job
        self.model_type = self.job.inputs["dataset_type"]

    def _build_dataset(self, document, dataset):
        document.add_heading(dataset["dataset_name"], 1)
        document.add_paragraph("TODO - add table")
        document.add_paragraph("TODO - build plot")

    def _build_model(self, document, dataset, model):
        document.add_heading(model["model_name"], 2)

        document.add_paragraph("TODO - build plot")

        # model options table
        if self.model_type == bmds.constants.CONTINUOUS:
            model_settings = {
                "BMR Type": bmrType[model["settings"]["bmrType"]],
                "BMRF": model["settings"]["bmr"],
                "Tail Proability": model["settings"]["tailProb"],
                "Confidence Level": model["settings"]["alpha"],
                "Distribution Type": model["settings"]["restriction"],
                "Variance Type": model["settings"]["varType"],
            }
        elif self.model_type == bmds.constants.DICHOTOMOUS:
            model_settings = {
                "Risk Type": riskType[model["settings"]["bmr_type"]],
                "BMRF": model["settings"]["bmr"],
                "Confidence Level": model["settings"]["alpha"],
            }
        else:
            raise ValueError("Unknown model_type")

        document.add_heading("Model options", level=3)
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        for key, value in model_settings.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)

        # dataset table
        model_data = {
            "Dependent Variable": dataset["column_names"]["doses"],
            "Number of Observation": len(dataset["doses"]),
        }
        if self.model_type == bmds.constants.DICHOTOMOUS:
            model_data["Independent Variable"] = dataset["column_names"]["incidences"]
        elif self.model_type == bmds.constants.CONTINUOUS:
            model_data["Independent Variable"] = dataset["column_names"]["responses"]
            model_data["Adverse Direction"] = dataset["adverse_direction"]

        document.add_heading("Model data", level=3)
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        for key, value in model_data.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)

        # benchmark dose table
        results = model["results"]
        benchmark_dose = {
            "BMD": results["bmd"],
            "BMDL": results["bmdl"],
            "BMDU": results["bmdu"],
            "AIC": results["aic"],
        }
        if self.model_type == bmds.constants.DICHOTOMOUS:
            benchmark_dose["P-value"] = results["gof"]["p_value"]
            benchmark_dose["Chi Square"] = results["gof"]["test_statistic"]
            benchmark_dose["D.O.F"] = results["gof"]["df"]
        document.add_heading("Benchmark Dose", level=3)
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        for key, value in benchmark_dose.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)

        # model parameter table
        params = {
            key: value
            for key, value in zip(results["fit"]["model"]["params"], results["fit"]["params"])
        }
        document.add_heading("Model Parameters", level=3)
        table = document.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "# of Parameters"
        table.rows[0].cells[1].text = str(len(params))
        table.rows[1].cells[0].text = "Variable"
        table.rows[1].cells[1].text = "Parameter"
        for key, value in params.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)

        # goodness of fit
        document.add_heading("Goodness of Fit", level=3)
        document.add_paragraph("TODO - add table")

        if self.model_type == bmds.constants.CONTINUOUS:
            # likelihood of interest
            document.add_heading("Likelihoods of Interest", level=3)
            document.add_paragraph("TODO - add table")

            # Test of Interest
            document.add_heading("Test of Interest", level=3)
            document.add_paragraph("TODO - add table")

        if self.model_type == bmds.constants.DICHOTOMOUS:
            document.add_heading("Analysis of Deviance", level=3)
            document.add_paragraph("TODO - add table")

        # CDF
        cdf = results["fit"]["bmd_dist"]
        document.add_heading("Cumulative Distribution Function (CDF)", level=3)
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.autofit
        heading_cells = table.rows[0].cells
        heading_cells[0].text = "Percentile"
        heading_cells[1].text = "BMD"
        for i in range(len(cdf[0])):
            row_cells = table.add_row().cells
            row_cells[0].text = str(cdf[1][i])
            row_cells[1].text = str(cdf[0][i])

        document.add_paragraph("TODO - build plot")

    def create_report(self) -> BytesIO:
        f = BytesIO()

        document = Document()
        document.add_heading(self.job.inputs.get("analysis_name", str(self.job.pk)), 0)

        if not self.job.is_finished:
            document.add_paragraph("Execution is incomplete; no report could be generated")
        elif self.job.has_errors:
            document.add_paragraph("Execution generated errors; no report can be generated")
        else:
            outputs: List[Dict] = self.job.outputs["outputs"]
            datasets: List[Dict] = self.job.inputs["datasets"]
            for output in outputs:
                dataset = datasets[output["dataset_index"]]
                self._build_dataset(document, dataset)
                for model in output["models"]:
                    self._build_model(document, dataset, model)

        document.save(f)
        return f


class ExportEngine:
    def __init__(self, job: Job):
        self.job = job
        self.model_type = self.job.inputs["dataset_type"]

    def _get_dataset_df(self) -> pd.DataFrame:
        """
        Return dataset and general analysis settings
        """
        created_time = self.job.created.strftime("%Y-%m-%d %H:%m:%S %Z")
        inputs = self.job.inputs

        if self.model_type == bmds.constants.CONTINUOUS:
            raise NotImplementedError()

        if self.model_type == bmds.constants.DICHOTOMOUS:
            datasets: List[Dict] = [
                dict(
                    analysis_name=inputs["analysis_name"],
                    analysis_description=inputs["analysis_description"],
                    dataset_type=inputs["dataset_type"],
                    bmds_version=inputs["bmds_version"] + "-ALPHA",
                    created=created_time,
                    dataset_index=i,
                    dataset_name=ds.get("dataset_name", ""),
                    doses=",".join([str(d) for d in ds["doses"]]),
                    ns=",".join([str(d) for d in ds["ns"]]),
                    incidences=",".join([str(d) for d in ds["incidences"]]),
                )
                for i, ds in enumerate(self.job.inputs["datasets"])
            ]
            return pd.DataFrame(datasets)

    def _get_model_df(self) -> pd.DataFrame:
        if self.model_type == bmds.constants.CONTINUOUS:
            raise NotImplementedError()

        data = []
        for output in self.job.outputs["outputs"]:
            for model in output["models"]:
                if not model["has_results"]:
                    # TODO handle this case
                    pass
                else:
                    results = model["results"]
                    data.append(
                        dict(
                            dataset_index=output["dataset_index"],
                            option_index=output["option_index"],
                            settings_bmr=model["settings"]["bmr"],
                            settings_bmr_type=bmrType[model["settings"]["bmr_type"]],
                            model_name=model["model_name"],
                            bmd=results["bmd"],
                            bmdl=results["bmdl"],
                            bmdu=results["bmdu"],
                            aic=results["aic"],
                        )
                    )

        return pd.DataFrame(data)

    def create_export(self) -> pd.DataFrame:
        # exit early if we don't have data for a report
        if not self.job.is_finished or self.job.has_errors:
            return pd.Series(
                data=["Job not finished or error occurred - cannot create report"], name="Status"
            ).to_frame()

        datasets = self._get_dataset_df()
        models = self._get_model_df()
        df = datasets.merge(models, on="dataset_index")

        return df
