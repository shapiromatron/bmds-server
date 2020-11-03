import io
import json

import pandas as pd
from django.core.exceptions import ValidationError
from docx import Document
from rest_framework import exceptions, mixins, status, viewsets
from rest_framework.authentication import TokenAuthentication
from rest_framework.decorators import action
from rest_framework.permissions import IsAdminUser
from rest_framework.response import Response

from . import models, renderers, serializers, tasks, validators


class JobViewset(mixins.CreateModelMixin, mixins.RetrieveModelMixin, viewsets.GenericViewSet):
    serializer_class = serializers.JobSerializer

    def not_ready_yet(self):
        content = "Outputs processing; not ready yet."
        return Response(content, status=status.HTTP_204_NO_CONTENT)

    @action(
        detail=True, methods=("patch",), url_path="patch-inputs",
    )
    def patch_inputs(self, request, *args, **kwargs):
        """
        Validate input and if successful, patch inputs on server side.
        """
        instance = self.get_object()
        data = request.data.get("data")
        edit_key = request.data.get("editKey", "")
        partial = bool(request.data.get("partial", False))

        # permission check
        if edit_key != instance.password:
            raise exceptions.PermissionDenied()

        if not isinstance(data, dict):
            raise exceptions.ValidationError("A `data` object is required")

        input_data = json.dumps(data)

        try:
            validators.validate_input(input_data, partial=partial)
        except ValidationError as err:
            raise exceptions.ValidationError(err.message)

        instance.inputs = input_data
        instance.save()

        serializer = self.get_serializer(instance)
        return Response(serializer.data)

    @action(detail=True, methods=("get",), renderer_classes=(renderers.TxtRenderer,))
    def inputs(self, request, *args, **kwargs):
        """
        Return inputs for selected job.
        """
        instance = self.get_object()
        fn = f"{instance.id}-inputs.json"
        resp = Response(instance.inputs)
        resp["Content-Disposition"] = f'attachment; filename="{fn}"'
        return resp

    @action(detail=True, methods=("post",))
    def execute(self, request, *args, **kwargs):
        """
        Attempt to execute the model.
        """
        instance = self.get_object()

        # permissions check
        if instance.password != request.data.get("editKey", ""):
            raise exceptions.PermissionDenied()

        # preflight execution check
        if not instance.inputs_valid():
            return Response("Invalid inputs", status_code=400)
        elif instance.is_executing:
            return Response("Execution already started", status_code=400)

        # start job execution
        instance.start_execute()

        instance.refresh_from_db()

        serializer = self.get_serializer(instance)
        return Response(serializer.data)

    @action(detail=True, methods=("get",), renderer_classes=(renderers.TxtRenderer,))
    def outputs(self, request, *args, **kwargs):
        """
        Return outputs for selected job
        """
        instance = self.get_object()

        if not instance.is_finished:
            return self.not_ready_yet()

        fn = f"{instance.id}-outputs.json"
        resp = Response(instance.outputs)
        resp["Content-Disposition"] = f'attachment; filename="{fn}"'
        return resp

    @action(detail=True, methods=("get",), renderer_classes=(renderers.XlsxRenderer,))
    def excel(self, request, *args, **kwargs):
        """
        Return Excel export of outputs for selected job
        """
        instance = self.get_object()

        # if not instance.is_finished:
        #     return self.not_ready_yet()
        # fn, wb = instance.get_excel()

        # TODO - set to temporary file for building UI - change later
        df = pd.DataFrame(data=dict(a=[1, 2, 3], b=[4, 5, 6]))
        f = io.BytesIO()
        df.to_excel(f, index=False)

        data = renderers.BinaryFile(data=f, filename=str(instance.id))
        return Response(data)

    @action(detail=True, methods=("get",), renderer_classes=(renderers.DocxRenderer,))
    def word(self, request, *args, **kwargs):
        """
        Return Word report for the selected job
        """
        instance = self.get_object()

        # if not instance.is_finished:
        #     return self.not_ready_yet()
        # fn, wb = instance.get_word()

        # TODO - set to temporary file for building UI - change later
        f = io.BytesIO()
        document = Document()
        document.add_heading("Continuous Result", 0)

        # created datasets for two different layouts data and data2.
        data = {
            "info": {
                "model": "Frequentist Exponential degree ",
                "Dataset Name": "Dataset Name1",
                "User notes": "User notes",
                "Dose-Response Model": "M[dose] = a * exp(±1 * b * dose)",
                "Variance Model": "Var[i]=alpha",
            },
            "model options": {
                "BMR Type": "Std. Dev.",
                "BMRF": "1",
                "Tail Probability": "-",
                "Confidence Level": "0.95",
                "Distribution TYpe": "Normal",
                "Variance": "Constant",
            },
            "model data": {
                "Dependent Variable": "Dose",
                "Independent Variables": "Mean",
                "Total # of observations": "4",
                "Adverse Direction": "Automatic",
            },
            "Benchmark Dose": {
                "BMD": "2.30",
                "BMDL": "0",
                "BMDU": "Infinity",
                "AIC": "56.40",
                "Test4 P-value": "0.53",
                "D.O.F": "1",
            },
            "Model Parameters": {
                "Number of Parameters": "3",
                "Variable": "Estimate",
                "a": "6.46",
                "b": "0.27",
                "log-alpha": "2.20",
            },
        }

        data2 = {
            "Goodness of Fit": {
                "Dose": [1, 4, 5, 5],
                "Size": [1, 4, 5, 5],
                "Estimated Median": [1, 4, 5, 5],
                "Cal'd Median": [1, 4, 5, 5],
                "observed Mean": [1, 4, 5, 5],
                "Estimated SD": [1, 4, 5, 5],
                "Calc'd SD": [1, 4, 5, 5],
                "Observed SD": [1, 4, 5, 5],
                "Scaled Residual": [1, 4, 5, 5],
            },
            "likelihood of Interest": {
                "Model": ["A1", "A1", "A3", "fitted", "R"],
                "Log Likelihood*": [-0.25, -22, -25, -25, -25],
                "# of parameters": [4, 6, 4, 3, 2],
                "AIC": [58.12, 56.06, 58.01, 56.402, 55.93],
            },
            "Test of Interest": {
                "Test": ["A1", "A1", "A3", "fitted", "R"],
                "2* Log": [-0.25, -22, -25, -25, -25],
                "Test df": [4, 6, 4, 3, 2],
                "P-value": [58.12, 56.06, 58.01, 56.402, 55.93],
            },
            "CDF": {"percentile": range(0, 100), "BMD": range(0, 100)},
        }

        # creating table for data layout
        for k, v in data.items():
            document.add_heading(k, level=1)
            table = document.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            table.autofit
            for name, desc in v.items():
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = k
                row_cells = table.add_row().cells
                row_cells[0].text = name
                row_cells[1].text = desc
        document.save(f)

        # creating table for data2 layout.
        for k, v in data2.items():
            document.add_heading(k, level=1)
            table = document.add_table(rows=1, cols=len(v))
            table.style = "Table Grid"
            table.autofit
            heading_cells = table.rows[0].cells
            keys_v = list(v.keys())
            for i in range(len(keys_v)):
                heading_cells[i].text = keys_v[i]
            list_values = list(v.values())
            res = [list(x) for x in zip(*list_values)]
            for i in res:
                str_i = [format(flt) for flt in i]
                cells = table.add_row().cells
                for e in range(len(str_i)):
                    cells[e].text = str_i[e]
        document.save(f)

        # plot is created and saved in projects root directory but not added to document.
        # todo add picture to document.
        plt.plot([0, 1, 2, 3, 4], [0, 3, 5, 9, 11])
        plt.xlabel("Months")
        plt.ylabel("Books Read")
        plt.savefig("picture.jpg")
        document.add_picture("picture.jpg", width=Inches(1.25))

        data = renderers.BinaryFile(data=f, filename=str(instance.id))
        return Response(data)

    def get_queryset(self):
        return models.Job.objects.all()


class DfileExecutorViewset(viewsets.ViewSet):

    permission_classes = (IsAdminUser,)
    authentication_classes = (TokenAuthentication,)

    def create(self, request):
        """
        Execute list of dfiles
        """
        payload = request.data.get("inputs", [])
        try:
            output = tasks.execute_dfile.delay(payload).get(timeout=120)
        except TimeoutError:
            output = {"timeout": True}
        return Response(output)
